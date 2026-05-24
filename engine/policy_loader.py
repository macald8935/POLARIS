"""Policy ingestion utilities for TXT, PDF, and DOCX files."""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def load_policy(file_path: str | Path) -> str:
    """Load a policy document and return normalized lower-case text."""

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Policy file not found: {path}")

    extension = path.suffix.lower()
    if extension == ".txt":
        text = _load_txt(path)
    elif extension == ".pdf":
        text = _load_pdf(path)
    elif extension == ".docx":
        text = _load_docx(path)
    else:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported policy format '{extension}'. Supported: {supported}")

    return normalize_policy_text(text)


def normalize_policy_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip().lower()


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _load_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(paragraphs + table_text)


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("Install pdfplumber to load PDF policy files.") from exc

    pages = []
    with pdfplumber.open(path) as pdf:
        raw_pages = [page.extract_text() or "" for page in pdf.pages]
    repeated = _detect_repeated_pdf_lines(raw_pages)

    for page_text in raw_pages:
        lines = []
        for line in page_text.splitlines():
            cleaned = line.strip()
            if not cleaned or cleaned in repeated or _looks_like_page_number(cleaned):
                continue
            lines.append(cleaned)
        pages.append("\n".join(lines))

    return "\n\n".join(page for page in pages if page.strip())


def _detect_repeated_pdf_lines(pages: list[str]) -> set[str]:
    if len(pages) < 3:
        return set()

    candidates = []
    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        candidates.extend(set(lines[:2] + lines[-2:]))

    counts = Counter(candidates)
    minimum_repetitions = max(2, int(len(pages) * 0.6))
    return {line for line, count in counts.items() if count >= minimum_repetitions}


def _looks_like_page_number(line: str) -> bool:
    return bool(re.fullmatch(r"(page\s*)?\d+(\s+of\s+\d+)?", line.lower()))
